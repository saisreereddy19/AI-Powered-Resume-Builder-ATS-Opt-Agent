
import os
import uuid
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from docx import Document
from dotenv import load_dotenv
import fitz  # PDF extraction
from google import genai
from pydantic import BaseModel
from typing import List, Dict

load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")
client = genai.Client(api_key=API_KEY)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

os.makedirs("storage", exist_ok=True)

# Utility to clean unwanted AI prompt intro text from enhanced resume
def clean_enhanced_resume(res_text):
    import re
    unwanted_phrases = [
        "okay, here’s a revised and enhanced version",
        "you are a world-class resume writer",
        "key improvements and explanations",
        "to use this resume",
        "••ats friendliness:••",
        "ATS-friendly, more detailed, and using impactful wording.",
        " enhanced version of your resume",
        "ATS-friendly,",
        " job description",
        "okay, here's the enhanced resume for",
    ]
    lines = res_text.lower().split('\n')
    core_start_idx = 0
    for idx, line in enumerate(lines):
        if line.strip() == '':
            continue
        if any(phrase.lower() in line for phrase in unwanted_phrases):
            continue
        else:
            core_start_idx = idx
            break
    orig_lines = res_text.split('\n')
    cleaned_lines = []
    skip_block = False
    for line in orig_lines[core_start_idx:]:
        if any(phrase.lower() in line.lower() for phrase in unwanted_phrases):
            skip_block = True
        if not skip_block and line.strip() != '':
            cleaned_lines.append(line)
    cleaned_text = "\n".join(cleaned_lines).strip()
    return cleaned_text

@app.post("/upload-pdf/")
async def upload_pdf(file: UploadFile = File(...)):
    try:
        pdf = fitz.open(stream=await file.read(), filetype="pdf")
        text = ""
        for page in pdf:
            text += page.get_text()
        return {"json": {"content": text}}
    except Exception as e:
        return {"error": str(e)}

@app.post("/score/")
def compute_score(req: dict):
    resume = req.get("resume_text", "").lower()
    jd = req.get("job_desc", "").lower()
    keyword_score = int((len(set(jd.split()) & set(resume.split())) / max(len(set(jd.split())), 1)) * 100)
    return {
        "semantic_score": keyword_score,
        "keyword_score": keyword_score,
        "final_score": keyword_score,
    }

@app.post("/enhance/")
def enhance_resume(req: dict):
    try:
        prompt = f"""
        Enhance the following resume based on the job description.
        Make it detailed, ATS-friendly, and improve wording.

        Resume:
        {req.get('resume_text', '')}

        Job Description:
        {req.get('job_desc', '')}
        """
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt
        )
        raw_text = response.text
        enhanced = clean_enhanced_resume(raw_text)
        return {"enhanced_text": enhanced}
    except Exception as e:
        return {"error": str(e)}

def generate_pdf(resume_text):
    file_id = str(uuid.uuid4())
    path = f"storage/{file_id}.pdf"
    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=50, leftMargin=50,
                            topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    Story = []

    for line in resume_text.split("\n"):
        line = line.strip()
        if not line:
            Story.append(Spacer(1, 0.1 * inch))
            continue
        if line.startswith("**") and line.endswith("**"):
            p = Paragraph(f"<b>{line[2:-2]}</b>", styles["Heading2"])
        else:
            p = Paragraph(line.replace("*", "•"), styles["Normal"])
        Story.append(p)
        Story.append(Spacer(1, 0.05 * inch))
    doc.build(Story)
    return file_id

@app.post("/download-pdf/")
def generate_pdf_route(data: dict):
    resume_text = clean_enhanced_resume(data.get("resume_text", ""))
    file_id = generate_pdf(resume_text)
    return {"file_id": file_id}

def build_docx(resume_text, out_path):
    doc = Document()
    for line in resume_text.split("\n"):
        line = line.strip()
        if line.startswith("**") and line.endswith("**"):
            doc.add_heading(line[2:-2], level=2)
        elif line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    doc.save(out_path)

@app.post("/download-docx/")
def generate_docx_route(data: dict):
    resume_text = clean_enhanced_resume(data.get("resume_text", ""))
    file_id = str(uuid.uuid4())
    path = f"storage/{file_id}.docx"
    build_docx(resume_text, path)
    return {"file_id": file_id}

class ResumeData(BaseModel):
    personal_info: Dict[str, str]
    education: List[str]
    skills: List[str]
    experience: List[str]
    projects: List[str]
    certifications: List[str]

def manual_resume_to_text(data: ResumeData) -> str:
    education_section = "- " + "\n- ".join(data.education)
    skills_section = "- " + "\n- ".join(data.skills)
    experience_section = "- " + "\n- ".join(data.experience)
    projects_section = "- " + "\n- ".join(data.projects)
    certifications_section = "- " + "\n- ".join(data.certifications)

    text = f"""Name: {data.personal_info.get('full_name', '')}
Email: {data.personal_info.get('email', '')}
Phone: {data.personal_info.get('phone', '')}
LinkedIn: {data.personal_info.get('linkedin', '')}
GitHub: {data.personal_info.get('github', '')}
Location: {data.personal_info.get('location', '')}

Education:
{education_section}

Skills:
{skills_section}

Experience:
{experience_section}

Projects:
{projects_section}

Certifications:
{certifications_section}
"""
    return text


@app.post("/manual-resume-text/")
def manual_resume_text(data: ResumeData):
    resume_text = manual_resume_to_text(data)
    return {"resume_text": resume_text}

app.mount("/storage", StaticFiles(directory="storage"), name="storage")
