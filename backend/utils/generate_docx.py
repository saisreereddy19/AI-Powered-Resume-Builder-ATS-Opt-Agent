from docx import Document

def build_docx(data, out_path):
    doc = Document()

    doc.add_heading(data["name"], level=1)
    doc.add_paragraph(data["email"])
    doc.add_paragraph(data["phone"])

    doc.add_heading("Experience", level=2)
    for exp in data["experience"]:
        doc.add_paragraph(exp)

    doc.save(out_path)
